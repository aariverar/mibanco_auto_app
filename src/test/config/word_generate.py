import os
import shutil
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
from PIL import ImageGrab
from PIL import Image
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service as ServiceChrome
from selenium import webdriver
import configparser
import platform
import psutil
from docxtpl import DocxTemplate
from allure import attach
import allure
# Constantes
PATH_RELATIVE_WORD = os.path.join(os.getcwd(), "src", "test", "resources","template","Plantilla2.docx")
TEMPLATE = 'template.png'
WORD_NAME_STANDAR = 'nombre_estandar_archivo_word.docx'
WORD_EXTENSION = '.docx'

def generate_word_from_table(nombre,correo,password):
    config = configparser.ConfigParser()
    config.read('mobile.properties')
    dispositivo = config.get('mobile', 'dispositivo')
    # Obtener el nombre del host
    hostname = platform.node()
    # Obtener el nombre de usuario
    username = psutil.users()[0].name
    # Obtener el sistema operativo y su versión
    os_info = platform.platform()
    # Obtener la información de la CPU
    cpu_info = platform.processor()
    # Obtener el número de núcleos de la CPU
    cpu_cores = psutil.cpu_count(logical=False)
    # Obtener la cantidad total de memoria RAM
    total_memory = psutil.virtual_memory().total
    # Obtener la zona horaria
    timezone_info = generar_hora()
    # Obtener la fecha y hora actual
    current_datetimex = generar_fecha()

    # Cargar la plantilla de Word
    doc = DocxTemplate(temp_word_file)
    # Renderizar la plantilla con los datos generados
    context = {
        'hostname': hostname,
        'username': username,
        'os_info': os_info,
        'cpu_info': cpu_info,
        'cpu_cores': cpu_cores,
        'total_memory': total_memory,
        'timezone_info': timezone_info,
        'current_datetimex': current_datetimex,
        'dispositivo': dispositivo,
        'nombre' : nombre,
        'correo': correo,
        'password': password
    }
    doc.render(context)
    doc.save(os.path.join(temp_word_file))

def start_up_word(name):
    insert_template = None
    try:
        
        # Crear una carpeta temporal para guardar el archivo de Word
        carpeta = 'Evidencias'
        os.makedirs(carpeta, exist_ok=True)

        global temp_word_file 
        temp_word_file = os.path.join(carpeta, f'{name}-{generar_secuencia()}{WORD_EXTENSION}')
        # Copiar la plantilla existente (si la hay) al directorio temporal
        file_unique = PATH_RELATIVE_WORD
        copy_existent_word(file_unique)
        # Crear un nuevo documento de Word
        global document
        document = Document(os.path.join(os.getcwd(), temp_word_file))
        
        # Agregar un párrafo al documento
        paragraph = document.add_paragraph()
        #run
        run = paragraph.add_run()
        # Guardar el documento de Word en un archivo
        
        document.save(os.path.join(os.getcwd(), temp_word_file))
        
        

    except Exception as e:
        print(f"[ERROR] {e}")
    finally:
        if insert_template:
            insert_template.close()

    print("[LOG] Word generado")

def copy_existent_word(file):
    try:
        # Copiar el archivo existente a la ubicación deseada
        copy_file = os.path.join(os.getcwd(), temp_word_file)
        shutil.copyfile(file, copy_file)
    except Exception as e:
        print(f"[ERROR] {e}")


def add_image_to_word(driver):
    try:
        document = Document(os.path.join(os.getcwd(), temp_word_file))
        # Agregar un párrafo al documento
        paragraph = document.add_paragraph()
        # Crear un nuevo objeto run para el párrafo
        run = paragraph.add_run()
        # Tomar una captura de pantalla
        if not os.path.exists('screenshots'):
            os.makedirs('screenshots')
        nombreImg= f"{generar_secuencia()}.png"
        filepath = os.path.join('screenshots', nombreImg)
        screenshot = driver.get_screenshot_as_png()
        allure.attach(driver.get_screenshot_as_png(), name=nombreImg, attachment_type=allure.attachment_type.PNG)
        with open(filepath, "wb") as file:
            file.write(screenshot)

        # Insertar la imagen en el documento Word
        #run.add_picture(filepath, width=Inches(6))
        picture = run.add_picture(filepath)
        picture.height = Cm(10)  # Altura de 10 cm
        picture.width = Cm(5)  # Ancho de 5 cm
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run.add_break()
        #os.remove(filepath)
        document.save(os.path.join(os.getcwd(), temp_word_file))
    except Exception as e:
        print(f"[ERROR] {e}")
    finally:
        print("[LOG] Imagen agregada a Word")

def send_text(texto):
    document = Document(os.path.join(os.getcwd(), temp_word_file))
    # Agregar un párrafo al documento
    paragraph = document.add_paragraph()

    # Crear un nuevo objeto run para el párrafo
    run = paragraph.add_run()

    run.text = f"Fecha : {generar_fecha()}, Hora : {generar_hora()} | {texto}"
    run.add_tab()
    run.font.name = "Century Gothic"
    run.font.size = Pt(9)
    document.save(os.path.join(os.getcwd(), temp_word_file))

def borar_word():
    try:
        os.remove(os.path.join(os.getcwd(), temp_word_file))
        print("[LOG] Word Eliminado")
    except Exception as e:
        print(f"[ERROR] {e}")

def end_to_word(status):
    try:
        # Renombrar el archivo con el estado final
        file_with_new_name = f"{temp_word_file.split('.docx')[0]}-{status.upper()}{WORD_EXTENSION}"
        os.rename(temp_word_file, file_with_new_name)

        print("[LOG] Word cerrado")

    except Exception as e:
        print(f"[ERROR] {e}")
    
def generar_secuencia():
    # Obtener la fecha y hora actual en el formato "dd-MM-yyyy_hh-mm-ss"
    now = datetime.datetime.now()
    formatted_date = now.strftime("%d-%m-%Y_%H-%M-%S")
    return formatted_date


def generar_fecha():
    # Obtener la fecha actual en el formato "dd/MM/yyyy"
    now = datetime.datetime.now()
    formatted_date = now.strftime("%d/%m/%Y")
    return formatted_date

def generar_hora():
    # Obtener la hora actual en el formato "hh:mm:ss"
    now = datetime.datetime.now()
    formatted_time = now.strftime("%H:%M:%S")
    return formatted_time