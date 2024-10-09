import os
import shutil
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from PIL import ImageGrab
from PIL import Image
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service as ServiceChrome
from selenium import webdriver
import configparser
import platform
import psutil
from docxtpl import DocxTemplate
import random

# Constantes
#PATH_RELATIVE_WORD = os.path.join(os.getcwd(), "src", "test", "resources","template","Plantilla2.docx")
PATH_RELATIVE_WORD = os.path.join(os.getcwd(), "src", "test", "resources","Plantilla.docx")
TEMPLATE = 'template.png'
WORD_NAME_STANDAR = 'nombre_estandar_archivo_word.docx'
WORD_EXTENSION = '.docx'

def generate_word_from_table(nombre,apellido,correo,telefono):
    config = configparser.ConfigParser()
    config.read('browser.properties')
    browser = config.get('General', 'browser')
    # Obtener el nombre del host
    hostname = platform.node()
    # Obtener el nombre de usuario
    username = psutil.users()[0].name
    # Obtener el sistema operativo y su versión
    os_info = platform.platform()
    # Obtener la información de la CPU
    cpu_info = platform.processor()
    # Obtener el número de núcleos de la CPU
    cpu_cores = psutil.cpu_count(logical=False)
    # Obtener la cantidad total de memoria RAM
    total_memory = psutil.virtual_memory().total
    # Obtener la zona horaria
    timezone_info = generar_hora()
    # Obtener la fecha y hora actual
    current_datetimex = generar_fecha()

    # Cargar la plantilla de Word
    doc = DocxTemplate(temp_word_file)
    # Renderizar la plantilla con los datos generados
    context = {
        'hostname': hostname,
        'username': username,
        'os_info': os_info,
        'cpu_info': cpu_info,
        'cpu_cores': cpu_cores,
        'total_memory': total_memory,
        'timezone_info': timezone_info,
        'current_datetimex': current_datetimex,
        'browser': browser,
        'nombre' : nombre,
        'apellido': apellido,
        'correo': correo,
        'telefono': telefono
    }
    doc.render(context)
    doc.save(os.path.join(temp_word_file))


def start_up_word(name):
    insert_template = None
    try:
        
        # Crear una carpeta temporal para guardar el archivo de Word
        carpeta = 'Evidencias'
        os.makedirs(carpeta, exist_ok=True)

        global temp_word_file 
        temp_word_file = os.path.join(carpeta, f'{name}-{generar_secuencia()}{WORD_EXTENSION}')
        # Copiar la plantilla existente (si la hay) al directorio temporal
        file_unique = PATH_RELATIVE_WORD
        copy_existent_word(file_unique)
        # Crear un nuevo documento de Word
        global document
        document = Document(os.path.join(os.getcwd(), temp_word_file))
        
        # Agregar un párrafo al documento
        #paragraph = document.add_paragraph()
        #run
        #run = paragraph.add_run()
        # Guardar el documento de Word en un archivo
        
        document.save(os.path.join(os.getcwd(), temp_word_file))
        
        

    except Exception as e:
        print(f"[ERROR] {e}")
    finally:
        if insert_template:
            insert_template.close()

    print("[LOG] Word generado")

def crear_tabla_inicio(casoPrueba,precondiciones,resultadoEsperado,context):
    # Añadir un párrafo al final del documento
    #doc.add_paragraph()
    ROW1=4
    COL1=4
    # Establecer la fuente y el tamaño del texto
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Añadir una tabla con 3 filas y 4 columnas
    table = document.add_table(rows=ROW1, cols=COL1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.style = 'Table Grid'
    #'Table Grid'
    # Ajustar el ancho de las columnas
    for column in table.columns:
        for cell in column.cells:
            cell.width = document.sections[0].page_width / 4.6

    # Añadir bordes y alinear verticalmente las celdas
    for i in range(ROW1):
        for o in range(COL1):
            table.cell(i, o).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


    # Rellenar la tabla con los datos
    set_cell_text(table.cell(0,0),"QA Agile: ",bold=True,background_color=colorRGB("GRIS"))
    set_cell_text(table.cell(0,1),"Equipo de automatización")

    set_cell_text(table.cell(0,2),"Fecha Ejecución: ",bold=True,background_color=colorRGB("GRIS"))
    set_cell_text(table.cell(0,3),generar_fecha())

    set_cell_text(table.cell(1,0),"Caso de Prueba: ",bold=True,background_color=colorRGB("GRIS"))
    set_cell_text(table.cell(1,1),casoPrueba)
    
    set_cell_text(table.cell(1,2),"Estado: ",bold=True,background_color=colorRGB("GRIS"))
    set_cell_text(table.cell(1,3),"")

    set_cell_text(table.cell(2,0),"Precondiciones:\n(Data,Accesos, ...) ",bold=True,background_color=colorRGB("GRIS"))
    set_cell_text(table.cell(2,1),precondiciones)

    set_cell_text(table.cell(2,2),"Entorno de prueba: ",bold=True,background_color=colorRGB("GRIS"))
    set_cell_text(table.cell(2,3),f"{context.deviceName} - Android {context.versionAndroid}")
    
    set_cell_text(table.cell(3,0),"Resultado Esperado: ",bold=True,background_color=colorRGB("GRIS"))
    table.cell(3, 1).merge(table.cell(3, 3))
    set_cell_text(table.cell(3,1),resultadoEsperado)


    # Añadir un párrafo al final del documento
    document.add_paragraph()
    document.save(os.path.join(os.getcwd(), temp_word_file))
    print("[LOG] Se crea tabla Inicio")

def copy_existent_word(file):
    try:
        # Copiar el archivo existente a la ubicación deseada
        copy_file = os.path.join(os.getcwd(), temp_word_file)
        shutil.copyfile(file, copy_file)
    except Exception as e:
        print(f"[ERROR] {e}")


def add_image_to_word(driver):
    try:
        document = Document(os.path.join(os.getcwd(), temp_word_file))
        # Agregar un párrafo al documento
        paragraph = document.add_paragraph()
        # Crear un nuevo objeto run para el párrafo
        run = paragraph.add_run()
        # Tomar una captura de pantalla
        if not os.path.exists('screenshots'):
            os.makedirs('screenshots')
        nombreImg= f"{generar_secuencia()}-{random.randint(10, 100)}.png"
        filepath = os.path.join('screenshots', nombreImg)
        screenshot = driver.get_screenshot_as_png()
        with open(filepath, "wb") as file:
            file.write(screenshot)

        # Insertar la imagen en el documento Word
        run.add_picture(filepath, width=Inches(2),height=Inches(3.5))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run.add_break()
        document.save(os.path.join(os.getcwd(), temp_word_file))
        return nombreImg
    except Exception as e:
        print(f"[ERROR] {e}")
    finally:
        print("[LOG] Imagen agregada a Word")

def add_image_to_word_web(driver):
    try:
        document = Document(os.path.join(os.getcwd(), temp_word_file))
        # Agregar un párrafo al documento
        paragraph = document.add_paragraph()
        # Crear un nuevo objeto run para el párrafo
        run = paragraph.add_run()
        # Tomar una captura de pantalla
        if not os.path.exists('screenshots'):
            os.makedirs('screenshots')
        nombreImg= f"{generar_secuencia()}-{random.randint(10, 100)}.png"
        filepath = os.path.join('screenshots', nombreImg)
        screenshot = driver.get_screenshot_as_png()
        with open(filepath, "wb") as file:
            file.write(screenshot)

        # Insertar la imagen en el documento Word
        run.add_picture(filepath, width=Inches(5),height=Inches(3))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run.add_break()
        document.save(os.path.join(os.getcwd(), temp_word_file))
        return nombreImg
    except Exception as e:
        print(f"[ERROR] {e}")
    finally:
        print("[LOG] Imagen agregada a Word")

def send_text(texto):
    print(f"[LOG]{texto}")
    document = Document(os.path.join(os.getcwd(), temp_word_file))
    # Agregar un párrafo al documento
    paragraph = document.add_paragraph()

    # Crear un nuevo objeto run para el párrafo
    run = paragraph.add_run()

    run.text = f"Fecha : {generar_fecha()}, Hora : {generar_hora()} | {texto}"
    run.add_tab()
    run.font.name = "Century Gothic"
    run.font.size = Pt(9)
    document.save(os.path.join(os.getcwd(), temp_word_file))

def borrar_word():
    try:
        os.remove(os.path.join(os.getcwd(), temp_word_file))
        print("[LOG] Word Eliminado")
    except Exception as e:
        print(f"[ERROR] {e}")

def end_to_word(status):
    try:
        #Actualizar el Estado de la prueba en el Word
        document = Document(os.path.join(os.getcwd(), temp_word_file))
        table = document.tables[0]
        set_cell_text(table.cell(1,3),status.upper(),background_color=colorStatus(status.upper()))
        document.save(os.path.join(temp_word_file))

        #Actualizar el nombre del Word
        file_with_new_name = f"{temp_word_file.split('.docx')[0]}-{status.upper()}{WORD_EXTENSION}"
        os.rename(temp_word_file, file_with_new_name)

        print("[LOG] Word cerrado")

    except Exception as e:
        print(f"[ERROR] {e}")
    
def generar_secuencia():
    # Obtener la fecha y hora actual en el formato "dd-MM-yyyy_hh-mm-ss"
    now = datetime.datetime.now()
    formatted_date = now.strftime("%d-%m-%Y_%H-%M-%S")
    return formatted_date

def set_cell_text(cell, text, bold=False, underline=False, background_color=None,alignment =WD_ALIGN_PARAGRAPH.CENTER,tipoLetra='Calibri',tamañoLetra=Pt(10)):
        run = cell.paragraphs[0].add_run(text)
        run.bold = bold
        run.underline = underline
        run.font.name = tipoLetra
        run.font.size = tamañoLetra
        cell.paragraphs[0].paragraph_format.alignment = alignment
        if background_color:
            cell_fill = OxmlElement('w:shd')
            cell_fill.set(qn('w:fill'),background_color)
            cell._element.get_or_add_tcPr().append(cell_fill)
        

def colorRGB(color):
    if color == "GRIS":
        return 'D3D3D3'
    else:
        return 'FAFAFA'
    
def colorStatus(status):
    if status == "PASSED":
        return '68FF00'
    else:
        return 'FF0000' 


def generar_fecha():
    # Obtener la fecha actual en el formato "dd/MM/yyyy"
    now = datetime.datetime.now()
    formatted_date = now.strftime("%d/%m/%Y")
    return formatted_date

def generar_hora():
    # Obtener la hora actual en el formato "hh:mm:ss"
    now = datetime.datetime.now()
    formatted_time = now.strftime("%H:%M:%S")
    return formatted_time